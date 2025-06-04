import re
import os
import shutil # Import shutil
from io import StringIO
from datetime import datetime
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
import docx # Added for python-docx interaction

CONFIG_FILE_PATH = "config.py"

def extract_text_from_pdf(pdf_path):
    """
    Extracts text content from a PDF file.
    The extracted text is returned as a single string.
    """
    output_string = StringIO()
    with open(pdf_path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        laparams = LAParams()
        device = TextConverter(rsrcmgr, output_string, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        print(f"Processing PDF: {pdf_path}")
        for i, page in enumerate(PDFPage.create_pages(doc)):
            print(f"  Processing page {i + 1}...")
            interpreter.process_page(page)
        print("PDF processing complete.")
    return output_string.getvalue()


# New function to check folder existence
def check_folder_exists(target_folder_path: str) -> bool:
    """Checks if a folder or file already exists at the given path."""
    return os.path.exists(target_folder_path)


def preprocess_text_initial(text):
    """
    Initial preprocessing: Removes DocuSign IDs, DigitalControls, pipe chars.
    DOES NOT NORMALIZE MULTIPLE SPACES TO ONE.
    """
    if text is None:
        return None
    text = re.sub(r"Docusign Envelope ID: [\w-]+\s*\f?", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\|[\w-]+DigitalControl_[\w_]+?\s*", " ", text)
    text = re.sub(r"\s+\|\s+", " ", text)
    text = re.sub(r"\|\s+", " ", text)
    text = re.sub(r"\s+\|", " ", text)
    return text.strip()



def normalize_multiple_spaces_in_text(text_segment):
    """Utility to reduce multiple spaces to one for text segments."""
    if text_segment is None:
        return None
    return re.sub(r"\s{2,}", " ", text_segment)



def preprocess_text_globally(text):
    """
    General preprocessing: Applies initial cleaning and then normalizes multiple spaces.
    This is a global cleaner used before detailed parsing.
    """
    if text is None:
        return None
    text = preprocess_text_initial(text)
    text = normalize_multiple_spaces_in_text(text)
    return text


def parse_any_legacy_contract_text(original_text_from_pdf: str) -> dict:
    """
    Parses the raw text extracted from a Legacy contract PDF to find specific data fields.
    Uses a series of regular expressions and text processing techniques.
    The `original_text_from_pdf` should be the direct output from `extract_text_from_pdf`.
    """
    data = {}

    # Preprocess text: initial cleaning for sensitive parsing (e.g., names before space normalization)
    # and global cleaning for general field extraction.
    text_for_sensitive_parsing = preprocess_text_initial(original_text_from_pdf)
    globally_cleaned_text = normalize_multiple_spaces_in_text(text_for_sensitive_parsing)

    def search_and_extract(pattern, text_to_search, group_index=1, default=None, flags=re.DOTALL | re.IGNORECASE):
        if not text_to_search:
            return default
        match = re.search(pattern, text_to_search, flags)
        if match:
            try:
                if group_index <= len(match.groups()) and match.group(group_index) is not None:
                    captured_value = match.group(group_index).strip()
                    captured_value = re.sub(r"Docusign Envelope ID:.*?\f", "", captured_value, flags=re.IGNORECASE | re.DOTALL).strip()
                    captured_value = re.sub(r"\|[\w-]+DigitalControl_[\w_]+?\s*", " ", captured_value).strip()
                    return re.sub(r"\s{2,}", " ", captured_value).strip()
                else:
                    return default
            except IndexError:
                return default
        return default

    def clean_phone(phone_str):
        if phone_str:
            return re.sub(r'\D', '', phone_str)
        return None

    def clean_currency(currency_str):
        if currency_str:
            return currency_str.replace('$', '').replace(',', '').strip()
        return None

    def normalize_state(state_str):
        if not state_str:
            return None
        state_str_upper = state_str.upper().strip()
        state_str_upper = re.sub(r'\s*\f\s*', '', state_str_upper)
        if "MISSISSIPPI" in state_str_upper or state_str_upper == "MS":
            return "MS"
        if "TENNESSEE" in state_str_upper or state_str_upper == "TN":
            return "TN"
        if "TEXAS" in state_str_upper or state_str_upper == "TX":
            return "TX"
        if len(state_str_upper) == 2:
            return state_str_upper
        return state_str.strip()[:2].upper()

    # --- Hardcoded Fields ---
    data['COUNTY'] = "DeSoto"
    data['SLR1ADR1'] = "5740 Getwell Road Building 8B"
    data['SLR1ADR2'] = "Southaven, MS 38672"
    data['AG701FRM'] = "Legacy Homes Realty, LLC"
    data['AG701LIC'] = "24125"
    data['AG701AD1'] = "5740 Getwell Rd Bldg 8B"
    data['AG701AD2'] = "Southaven, MS 38672"
    data['AG701PH'] = "6629322282"
    # --- End Hardcoded Fields ---

    data['INCITY'] = 'X'
    data['INCOUNTY'] = 'X'
    data['DEPHELD'] = 'Seller'
    data['POSSION'] = 'Fee Simple'
    data['UNDNAME'] = 'Chicago Title Insurance Company'
    data['SLR1NAM1'] = search_and_extract(r"Parties\s*-\s*(LEGACY NEW HOMES, LLC.*?)\s*hereafter called SELLER", globally_cleaned_text)
    data['SLR1REL1'], data['SLR1NAM2'] = '', None
    data['SLR1CELL1'], data['SLR1EMAIL'], data['SLR1CELL2'], data['SLR1EMAIL2'] = None, None, None, None
    data['PARCELID'] = None
    data['PLISTINGAGENT'], data['MTDTTYPE'] = '2%', 'Deed of Trust'

    data['SALEPRIC'] = clean_currency(search_and_extract(r"Full Purchase Price\s*\$?([\d,]+\.\d{2})", globally_cleaned_text))
    data['DEPOSIT'] = clean_currency(search_and_extract(r"Deposit held by\s*LEGACY NEW HOMES,LLC\s*\$?([\d,]+\.\d{2})", globally_cleaned_text))
    if not data['DEPOSIT']:
        data['DEPOSIT'] = clean_currency(search_and_extract(r"DEPOSIT Held by Legacy New Homes, LLC\s*\$?([\d,]+\.\d{2})", globally_cleaned_text))

    # --- START: SETTDATE Extraction Logic (Rewritten with user's reliable approach) ---
    data['SETTDATE'] = None
    anchor = "Home is to close on or before"
    idx = original_text_from_pdf.find(anchor)
    if idx != -1:
        after = original_text_from_pdf[idx + len(anchor):]
        for token in re.split(r"\s+", after):            # walk tokens after anchor
            if token.count('/') > 4:                     # first token with >4 “/” ends block
                # keep only digits, slashes, and spaces
                filtered = ''.join(
                    ch if (ch.isdigit() or ch == '/' or ch == ' ') else ' '
                    for ch in token
                )
                # grab **all** date patterns in the filtered string
                dates = re.findall(r"\d{1,2}/\d{1,2}/\d{4}", filtered)
                if dates:
                    def _p(d):
                        try:
                            return datetime.strptime(d, "%m/%d/%Y")
                        except ValueError:
                            return datetime.min
                    data['SETTDATE'] = max(dates, key=_p)   # <-- latest date wins
                break
    # --- END: SETTDATE Extraction Logic ---

    byr1_name, byr2_name, byr1_rel = None, None, ''
    title_block_match = re.search(r"wishes to take title as follows:\s*(.*?)(?=\s*Please List whether BUYER is:|\s*Single Person|\s*Married Person|\s*Investor|$)", text_for_sensitive_parsing, re.IGNORECASE | re.DOTALL)
    if title_block_match:
        raw_names_str = title_block_match.group(1).strip()
        cleaned_names_str = raw_names_str
        if cleaned_names_str.lower().endswith("please"):
            match_please_end = re.match(r"(.*?)(\w*Please)$", cleaned_names_str, re.IGNORECASE)
            if match_please_end:
                word_before_please = match_please_end.group(2)[:-6] if len(match_please_end.group(2)) >= 6 else ""
                cleaned_names_str = (match_please_end.group(1) + word_before_please).strip()
        if " and " in cleaned_names_str.lower():
            parts = re.split(r"\s+and\s+", cleaned_names_str, maxsplit=1, flags=re.IGNORECASE)
            byr1_name = parts[0].strip()
            if len(parts) > 1 and parts[1].strip():
                byr2_name = parts[1].strip()
                byr1_rel = 'and'
        elif re.search(r"\s{3,}", cleaned_names_str):
            parts = re.split(r"\s{3,}", cleaned_names_str, maxsplit=1)
            byr1_name = parts[0].strip()
            if len(parts) > 1 and parts[1].strip():
                byr2_name = parts[1].strip()
                byr1_rel = 'and'
        else:
            byr1_name = cleaned_names_str.strip()
    data['BYR1NAM1'], data['BYR1NAM2'], data['BYR1REL1'] = byr1_name, byr2_name, byr1_rel

    buyer_contact_block_match = re.search(r"hereafter called BUYER\(s\), whose address, phone numbers, and email address(?:es)? are listed below\s*(.*?)\s*hereby agree", globally_cleaned_text, re.DOTALL | re.IGNORECASE)
    if buyer_contact_block_match:
        contact_details_str = buyer_contact_block_match.group(1)
        buyer_contacts_pattern = r"(\d+[\w\s\.,#-]*?(?:Street|St|Road|Rd|Drive|Dr|Avenue|Ave|Lane|Ln|Cove|Cv|Court|Ct|Place|Pl|Boulevard|Blvd))\s+([A-Za-z\s'-]+?)\s+(MS|TN|TX|MISSISSIPPI|TENNESSEE|TEXAS)\s+(\d{5})\s*\(?(\d{3})\)?\s*(\d{3}-\d{4})\s+([\w\.@-]+)"
        buyer_contacts_fallback_pattern = r"(\d+[\w\s\.,#-]*?\s\w+)\s+([A-Za-z\s'-]+?)\s+(MS|TN|TX|MISSISSIPPI|TENNESSEE|TEXAS)\s+(\d{5})\s*\(?(\d{3})\)?\s*(\d{3}-\d{4})\s+([\w\.@-]+)"
        buyer_contacts = re.findall(buyer_contacts_pattern, contact_details_str, re.IGNORECASE)
        if not buyer_contacts:
            buyer_contacts = re.findall(buyer_contacts_fallback_pattern, contact_details_str, re.IGNORECASE)
        if len(buyer_contacts) > 0:
            b1 = buyer_contacts[0]
            data['BYR1ADR1'] = b1[0].strip()
            data['BYR1ADR2'] = f"{b1[1].strip()}, {normalize_state(b1[2])} {b1[3].strip()}"
            data['BYR1CELL1'] = clean_phone(f"{b1[4]}{b1[5]}")
            data['BYR1EMAIL'] = b1[6].strip()
        if len(buyer_contacts) > 1:
            b2 = buyer_contacts[1]
            data['BYR1CELL2'] = clean_phone(f"{b2[4]}{b2[5]}")
            data['BYR1EMAIL2'] = b2[6].strip()
    for k_buyer_contact in ['BYR1ADR1', 'BYR1ADR2', 'BYR1CELL1', 'BYR1EMAIL', 'BYR1CELL2', 'BYR1EMAIL2']:
        data.setdefault(k_buyer_contact, None)

    prop_match = re.search(r"Lot\s+([\w\d]+)(?:\s*Plan/Elevation\s+[\w\s\d.-]+?)?\s*Subdivision\s+([\w\s\d.-]+?Phase\s*\d+|[\w\s\d.-]+?Section\s*\w+\s*Phase\s*\d+|[\w\s\d.-]+?)\s*Address\s+([\d\w\s.-]+?(?:Lane|Drive|Road|Cove|Street|St|Ave|Dr))\s+([A-Za-z\s'-]+?)\s+(MISSISSIPPI|MS|TENNESSEE|TN)\s+(\d{5})", globally_cleaned_text, re.IGNORECASE | re.DOTALL)
    if prop_match:
        data['LORU'], data['LOTUNIT'] = "Lot", prop_match.group(1).strip()
        data['SUBDIVN'] = prop_match.group(2).strip()
        data['PROPSTRE'] = prop_match.group(3).strip()
        data['PROPCITY'] = prop_match.group(4).strip()
        data['STATELET'] = normalize_state(prop_match.group(5))
        data['PROPZIP'] = prop_match.group(6).strip()
    else:
        prop_match_fallback = re.search(r"Lot\s+([\w\d]+).*?Subdivision\s+(.*?)\s*Address\s+([\d\w\s.-]+)\s+([A-Za-z\s'-]+?)\s+(MISSISSIPPI|MS|TENNESSEE|TN)\s+(\d{5})", globally_cleaned_text, re.IGNORECASE | re.DOTALL)
        if prop_match_fallback:
            data['LORU'], data['LOTUNIT'] = "Lot", prop_match_fallback.group(1).strip()
            data['SUBDIVN'] = prop_match_fallback.group(2).strip()
            full_addr = prop_match_fallback.group(3).strip()
            city_prop = prop_match_fallback.group(4).strip()
            if city_prop.lower() in full_addr.lower() and full_addr.lower() != city_prop.lower():
                parts = full_addr.rsplit(city_prop, 1)
                data['PROPSTRE'] = parts[0].strip() if len(parts) > 1 and parts[0].strip() else full_addr
            else:
                data['PROPSTRE'] = full_addr
            data['PROPCITY'] = city_prop
            data['STATELET'] = normalize_state(prop_match_fallback.group(5))
            data['PROPZIP'] = prop_match_fallback.group(6).strip()
        else:
            data['LORU'], data['LOTUNIT'], data['SUBDIVN'], data['PROPSTRE'], data['PROPCITY'], data['STATELET'], data['PROPZIP'] = (
                None, None, None, None, None, None, None)

    agent_section_overall_match = re.search(r"AGENCY DISCLOSURE\s*-\s*\(check one\):(.*?)(?=13\.\s*ARBITRATION|14\.\s*ARBITRATION|SIGNING BELOW|BY SIGNING BELOW|Property Condition Disclosure|$)", globally_cleaned_text, re.DOTALL | re.IGNORECASE)
    agent_text_block_overall = agent_section_overall_match.group(1).strip() if agent_section_overall_match else globally_cleaned_text
    listing_agent_block_match = re.search(r"Listing Agency\s*(.*?)(?=Selling Agency|$)", agent_text_block_overall, re.DOTALL | re.IGNORECASE)
    listing_agent_block_text = listing_agent_block_match.group(1).strip() if listing_agent_block_match else ""
    if listing_agent_block_text:
        data['AG701NAM'] = search_and_extract(r"Listing Agent\s+([\w\s,-]+?)(?=\s*Business Phone)", listing_agent_block_text)
        if data['AG701NAM']:
            data['AG701NAM'] = data['AG701NAM'].replace(',', '').strip()
        l_agent_phone_match = re.search(r"Listing Agent\s+[\w\s,-]+?Business Phone\s*([()\d\s-]+?)(?=\s*Address|\s*Email)", listing_agent_block_text, re.IGNORECASE)
        data['AG701MO'] = clean_phone(l_agent_phone_match.group(1)) if l_agent_phone_match else None
        data['AG701EMAIL'] = search_and_extract(r"Email\s+([\w\.@-]+?)(?=\s+License #:\s*Agent|\s*$)", listing_agent_block_text)
        data['AG701CONTLIC'] = search_and_extract(r"License #:\s*Agent\s*(.*?)(?=Selling Agency|$)", listing_agent_block_text)
    else:
        for k in ['AG701NAM', 'AG701MO', 'AG701EMAIL', 'AG701CONTLIC']:
            data.setdefault(k, None)

    selling_agent_block_match = re.search(r"Selling Agency\s*(.*?)(?=13\.\s*ARBITRATION|14\.\s*ARBITRATION|Revised\s+\d{2}/\d{2}/\d{2}|$)", agent_text_block_overall, re.DOTALL | re.IGNORECASE)
    selling_agent_block_text_for_details = selling_agent_block_match.group(1).strip() if selling_agent_block_match else ""
    if selling_agent_block_text_for_details:
        data['AG702FRM'] = search_and_extract(r"^([\w\s.,'&@#-]+?)(?=\s*Selling Agent|\s*Business Phone)", selling_agent_block_text_for_details)
        data['AG702LIC'] = search_and_extract(r"License #:\s*Firm\s*([S\d][\w-]+?)(?=\s*License #:\s*Agent|\s*Email:|$)", selling_agent_block_text_for_details)
        if data['AG702LIC']: data['AG702LIC'] = data['AG702LIC'].replace(" ", "")
        s_addr_match = re.search(r"Address:\s*([\w\s\.\d#-]+(?:Street|Parkway|Road|Rd)?(?:,\s*\#?\w+)?)\s*,\s*([A-Za-z\s'-]+?)\s*,\s*(MISSISSIPPI|MS|TENNESSEE|TN|TEXAS|TX)(?:,\s*(\d{5})(?:,\s*United States of America)?)?", selling_agent_block_text_for_details, re.IGNORECASE)
        zip_s = None
        if s_addr_match:
            data['AG702AD1'] = s_addr_match.group(1).strip(); city_s = s_addr_match.group(2).strip().replace(',', ''); state_s = normalize_state(s_addr_match.group(3))
            if s_addr_match.group(4): zip_s = s_addr_match.group(4).strip()
            else:
                zip_s_alt_match = re.search(r"(?:MISSISSIPPI|MS|TENNESSEE|TN|TEXAS|TX)(?:,\s*United States of America)?\s*(\d{5})", selling_agent_block_text_for_details, re.IGNORECASE)
                if zip_s_alt_match: zip_s = zip_s_alt_match.group(1).strip()
            data['AG702AD2'] = f"{city_s}, {state_s} {zip_s}" if zip_s else f"{city_s}, {state_s}"
        else: data['AG702AD1'], data['AG702AD2'] = None, None
        data['AG702PH'] = clean_phone(search_and_extract(r"Business Phone\s*([()\d\s-]+?)(?=\s+Address:|\s+Selling Agent)", selling_agent_block_text_for_details))
        data['AG702NAM'] = search_and_extract(r"Selling Agent\s+([\w\s,-]+?)(?=\s*Business Phone|,Business Phone)", selling_agent_block_text_for_details)
        if data['AG702NAM']: data['AG702NAM'] = data['AG702NAM'].replace(',', '').strip()
        sa_phone_match = re.search(r"Selling Agent\s+[\w\s,-]+?Business Phone\s*([()\d\s-]+?)(?=\s*Address|\s*Email)", selling_agent_block_text_for_details, re.IGNORECASE)
        data['AG702MO'] = clean_phone(sa_phone_match.group(1)) if sa_phone_match else None
        data['AG702EMAIL'] = search_and_extract(r"Email:\s*([\w\.@-]+)", selling_agent_block_text_for_details)
    else:
        for k in ['AG702FRM', 'AG702PH', 'AG702AD1', 'AG702AD2', 'AG702LIC', 'AG702NAM', 'AG702MO', 'AG702EMAIL']: data.setdefault(k, None)

    ag702contlic_pattern_full_text = r"Selling Agency.*?License #:\s*Firm\s*(?:[S]-)?\d[\w-]*\s*License #:\s*Agent\s*(.{1,20}?)(?=\s*\d{1,2}\.\s*ARBITRATION)"
    data['AG702CONTLIC'] = search_and_extract(ag702contlic_pattern_full_text, globally_cleaned_text)
    if not data.get('AG702CONTLIC'):
         data['AG702CONTLIC'] = search_and_extract(r"Selling Agency.*?License #:\s*Agent\s*(.{1,20}?)(?=\s*\d{1,2}\.\s*ARBITRATION)", globally_cleaned_text)
    data.setdefault('AG702CONTLIC', None)

    listing_agent_comm_pct = 2.0
    buyer_agent_comm_match = re.search(r"buyer’s agent compensation of\s+(\d+)%", globally_cleaned_text, re.IGNORECASE)
    if buyer_agent_comm_match:
        try: data['COMPCT'] = f"{listing_agent_comm_pct + float(buyer_agent_comm_match.group(1))}%"
        except ValueError: data['COMPCT'] = f"{listing_agent_comm_pct}%"
    else: data['COMPCT'] = f"{listing_agent_comm_pct}%"
    
    return data

def get_all_legacy_contract_field_names() -> list[str]:
    """
    Returns a comprehensive list of all possible field names (keys) that
    the `parse_any_legacy_contract_text` function can extract and include
    in its returned dictionary. This list is based on a review of the
    keys assigned in that function.
    """
    # This list should be manually maintained and updated if
    # `parse_any_legacy_contract_text` changes which keys it can return.
    return [
        "COUNTY", "SLR1ADR1", "SLR1ADR2", "AG701FRM", "AG701LIC", "AG701AD1",
        "AG701AD2", "AG701PH", "INCITY", "INCOUNTY", "DEPHELD", "POSSION",
        "UNDNAME", "SLR1NAM1", "SLR1REL1", "SLR1NAM2", "SLR1CELL1", "SLR1EMAIL",
        "SLR1CELL2", "SLR1EMAIL2", "PARCELID", "PLISTINGAGENT", "MTDTTYPE",
        "SALEPRIC", "DEPOSIT", "SETTDATE", "BYR1NAM1", "BYR1NAM2", "BYR1REL1",
        "BYR1ADR1", "BYR1ADR2", "BYR1CELL1", "BYR1EMAIL", "BYR1CELL2", "BYR1EMAIL2",
        "LORU", "LOTUNIT", "SUBDIVN", "PROPSTRE", "PROPCITY", "STATELET", "PROPZIP",
        "AG701NAM", "AG701MO", "AG701EMAIL", "AG701CONTLIC", "AG702FRM", "AG702LIC",
        "AG702AD1", "AG702AD2", "AG702PH", "AG702NAM", "AG702MO", "AG702EMAIL",
        "AG702CONTLIC", "COMPCT"
    ]

def format_name(name1_str, name2_str=None):
    """
    Formats one or two names into a standardized string.
    Handles company names and individual names with different formats.
    """
    COMPANY_KEYWORDS = ["home", "prop", "llc", "inc", "custom", "build", "dev"]

    if any(keyword in name1_str.lower() for keyword in COMPANY_KEYWORDS if keyword): # ensure keyword is not empty
        return name1_str.upper()

    def _parse_name(name_str):
        if not name_str or not name_str.strip():
            return None, None
        
        name_str = name_str.strip()
        
        if ',' in name_str:
            parts = name_str.split(',', 1)
            last = parts[0].strip()
            first = parts[1].strip() if len(parts) > 1 else None
            return first, last
        else:
            parts = name_str.split()
            if not parts: # Should be caught by initial strip and check, but as safeguard
                return None, None
            if len(parts) == 1:
                return None, parts[0].strip() # Treat as last name
            else:
                last = parts[-1].strip()
                first = " ".join(parts[:-1]).strip()
                return first, last

    first1, last1 = _parse_name(name1_str)

    if not last1: # If name1_str was empty or unparseable to a last name
        return "" # Or a sensible default like "UNKNOWN NAME"

    if not name2_str or not name2_str.strip():
        # Only name1 is provided or name2 is empty
        if first1:
            return f"{last1.upper()}, {first1}"
        else:
            return last1.upper()

    first2, last2 = _parse_name(name2_str)

    if not last2: # If name2_str was empty or unparseable to a last name
        # Treat as if name2_str was not provided
        if first1:
            return f"{last1.upper()}, {first1}"
        else:
            return last1.upper()

    # Both names are provided and parseable
    if last1.upper() == last2.upper():
        return f"{last1.upper()}, {first1} & {first2}"
    else:
        # Ensure first names are not None before joining
        f1 = first1 if first1 else ""
        f2 = first2 if first2 else ""
        # Construct combined first names, handling cases where one or both might be empty
        combined_first_names = f"{f1} {f2}".strip() 
        if combined_first_names:
             return f"{last1.upper()} & {last2.upper()}, {combined_first_names}"
        else: # Both first names were None or empty
             return f"{last1.upper()} & {last2.upper()}"


def generate_legacy_folder_name(extracted_data: dict) -> str:
    """
    Generates a standardized folder name for a Legacy contract based on extracted data,
    primarily using the first buyer's name (BYR1NAM1).
    The format is typically "LastName, FirstName  25-" or "CompanyName  25-".
    """
    byr1nam1 = extracted_data.get('BYR1NAM1')
    if byr1nam1:
        name_parts = byr1nam1.split(',')
        if len(name_parts) > 1:
            last_name = name_parts[0].strip()
            first_name = name_parts[1].strip()
        else:
            name_parts_space = byr1nam1.split()
            if len(name_parts_space) > 1:
                last_name = name_parts_space[-1]
                first_name = " ".join(name_parts_space[:-1])
            else:
                last_name = byr1nam1
                first_name = ""
        first_name_safe = re.sub(r'[<>:"/\\|?*]', '', first_name)
        last_name_safe = re.sub(r'[<>:"/\\|?*]', '', last_name)
        return f"{last_name_safe}  25-"
    return "Unknown_Entity  25-"


# New function to copy PDF
def copy_pdf_to_folder(source_pdf_path: str, destination_folder_path: str, pdf_filename: str) -> tuple[bool, str]:
    """Copies the source PDF to the destination folder."""
    full_dest_pdf_path = os.path.join(destination_folder_path, pdf_filename)
    try:
        shutil.copy2(source_pdf_path, full_dest_pdf_path)
        print(f"INFO: Successfully copied {pdf_filename} to {destination_folder_path}")
        return True, f"Successfully copied {pdf_filename} to {destination_folder_path}"
    except (IOError, shutil.Error) as e:
        print(f"ERROR: Error copying {pdf_filename}: {e}")
        return False, f"Error copying {pdf_filename}: {e}"


# New function to get initial folder name and data for GUI checks
def get_initial_legacy_folder_name_and_data(pdf_file_path: str) -> tuple[str | None, dict | None, str | None]:
    """
    Extracts text, parses it, generates a folder name, and returns these.
    This function serves as a preliminary step, often called by the GUI, 
    to get essential information for checks like folder existence before 
    committing to the full folder creation and file population process.

    Args:
        pdf_file_path: The path to the PDF file to be processed.

    Returns:
        A tuple containing:
            - generated_name (str | None): The proposed folder name.
            - extracted_data (dict | None): The data extracted from the PDF.
            - error_message (str | None): An error message if any issue occurred, otherwise None.
    """
    try:
        print(f"INFO: Initial PDF text extraction for folder name generation: {pdf_file_path}")
        raw_text = extract_text_from_pdf(pdf_file_path)
        if not raw_text:
            error_msg = f"Could not extract any text from {pdf_file_path} for folder naming."
            print(f"ERROR: {error_msg}")
            return None, None, error_msg
        
        print(f"INFO: Initial parsing for folder name generation: {pdf_file_path}")
        extracted_data = parse_any_legacy_contract_text(raw_text)
        # Check for essential data needed for folder name (e.g., BYR1NAM1)
        if not extracted_data or not extracted_data.get('BYR1NAM1'):
            error_msg = f"Failed to parse essential data (like BYR1NAM1) from {pdf_file_path} for folder naming."
            print(f"ERROR: {error_msg}")
            return None, extracted_data, error_msg # Return partial data if any for context

        folder_name = generate_legacy_folder_name(extracted_data)
        return folder_name, extracted_data, None # Success
    except Exception as e:
        error_msg = f"An error occurred during initial processing for {pdf_file_path}: {e}"
        print(f"CRITICAL ERROR: {error_msg}")
        return None, None, error_msg


def create_legacy_contract_folder_structure(final_folder_path: str, extracted_data: dict) -> tuple[str, bool]:
    """
    Creates the specific folder structure for a Legacy contract at the `final_folder_path`.
    This includes:
    - The main client folder.
    - An "overlay.pxt" file containing all extracted data.
    - A "Setup" subfolder.
    - A "TitleSearch" subfolder.
    - "file label.docx" and "setupdocs.docx" (placeholder content) within the "Setup" subfolder.

    Args:
        final_folder_path: The absolute path where the folder structure should be created.
                           This path is determined by the GUI, possibly after user input for renaming.
        extracted_data: A dictionary containing data extracted from the PDF, used to populate
                        the .pxt file and potentially other generated files.

    Returns:
        A tuple containing:
            - final_folder_path (str): The path where the structure was attempted.
            - success (bool): True if creation was successful, False otherwise.
    """
    # final_folder_path is the full absolute path, decided by the GUI after any negotiations (e.g., renaming).
    print(f"INFO: Creating Legacy contract folder at: {final_folder_path}")

    try:
        os.makedirs(final_folder_path, exist_ok=True)

        # Create "overlay.pxt" with all extracted data
        entity_list_content = []
        if extracted_data: # Ensure extracted_data is not None
            for key, value in extracted_data.items():
                entity_list_content.append(f"{key}= {value}")
        
        with open(os.path.join(final_folder_path, "overlay.pxt"), "w") as f:
            f.write("\n".join(entity_list_content))
            f.write("\n\n--- End of Extracted Data ---")

        # Placeholder files
        # Use final_folder_path's basename for user-facing messages/content if needed
        # folder_basename = os.path.basename(final_folder_path) # No longer directly used for simple label
        
        # Create subfolders first (ensure "Setup" exists before writing files into it)
        setup_subfolder_path = os.path.join(final_folder_path, "Setup")
        os.makedirs(setup_subfolder_path, exist_ok=True)
        os.makedirs(os.path.join(final_folder_path, "TitleSearch"), exist_ok=True)

        # Remove old placeholder label creation in "Setup"
        # file_label_path = os.path.join(setup_subfolder_path, "file label.docx")
        # with open(file_label_path, "w") as f:
        #     f.write(f"File Label for: {folder_basename}\n")
        #     if extracted_data:
        #         f.write(f"Property Address: {extracted_data.get('PROPSTRE', 'N/A')}\n")
        
        # Create setupdocs.docx in "Setup" subfolder
        setup_docs_path = os.path.join(setup_subfolder_path, "setupdocs.docx")
        with open(setup_docs_path, "w") as f:
            f.write(f"Setup Documents for: {os.path.basename(final_folder_path)}\n") # Use basename here

        # Generate the new Label.docx in the "Setup" subfolder
        template_label_path = "templates/Label.docx" # Assuming this path is correct relative to execution
        # Ensure setup_subfolder_path is defined before this line (it is, a few lines above)
        output_label_path = os.path.join(setup_subfolder_path, "Label.docx")
        
        if extracted_data: # Ensure there's data for the label
            label_index = get_next_label_index()
            label_generated = generate_label_docx(template_label_path, output_label_path, extracted_data, label_index)
            if label_generated:
                update_label_index(label_index)
                print(f"INFO: Successfully generated and updated label index for {output_label_path}")
            else:
                print(f"WARNING: Failed to generate label for {output_label_path}")
        else:
            print(f"WARNING: No extracted_data available, skipping label generation for {final_folder_path}")

        print(f"SUCCESS: Created folder structure in '{final_folder_path}'.")
        return final_folder_path, True # Return path and success
    except OSError as e:
        print(f"ERROR: Could not create folder structure in '{final_folder_path}'. Error: {e}")
        return final_folder_path, False # Return path and failure


def handle_legacy_contract_processing(
    pdf_file_paths, # Though we might only use the first one if GUI passes one by one
    user_selected_output_dir: str, 
    processed_folder_name: str, # This is the name decided by GUI (original or renamed)
    extracted_data_from_gui: dict # This is the data from get_initial_legacy_folder_name_and_data
    ):
    """
    Main handler for the core logic of processing "Legacy" contracts.
    This function is called by the GUI after:
    1. Initial data extraction and folder name generation (`get_initial_legacy_folder_name_and_data`).
    2. User validation/confirmation of the output folder name (including overwrite/rename dialogs).
    
    It assumes the `processed_folder_name` and `extracted_data_from_gui` are finalized and correct.
    Its main responsibility is to call `create_legacy_contract_folder_structure`.
    The actual PDF text extraction and parsing are expected to have happened in the
    `get_initial_legacy_folder_name_and_data` step, and the results passed via `extracted_data_from_gui`.

    Args:
        pdf_file_paths: A list of PDF file paths (though typically only the first is used for Legacy).
        user_selected_output_dir: The base directory selected by the user for output.
        processed_folder_name: The final, confirmed name for the client-specific folder.
        extracted_data_from_gui: The dictionary of data extracted by `get_initial_legacy_folder_name_and_data`.

    Returns:
        A tuple containing:
            - created_folder_path (str | None): The full path to the created folder if successful, else None.
            - message_string (str): A message detailing the outcome of the operation.
    """
    # Input validation (should ideally be guaranteed by GUI calling sequence)
    if not pdf_file_paths:
        return None, "No PDF files provided for processing."
    if not user_selected_output_dir: # Added for robustness
        return None, "Output directory not specified."
    if not processed_folder_name:
        return None, "No folder name provided for processing."
    if not extracted_data_from_gui:
        return None, "No extracted data provided for processing."

    # Construct the final, absolute path for the client-specific folder
    final_folder_path = os.path.join(user_selected_output_dir, processed_folder_name)
    
    # Delegate the actual folder and file creation
    created_path, success = create_legacy_contract_folder_structure(
        final_folder_path, 
        extracted_data_from_gui
    )

    if success:
        # The GUI already has the extracted_data, so we just return the path and success message.
        return created_path, f"Successfully created folder structure in '{processed_folder_name}'"
    else:
        return None, f"Failed to create folder structure for '{processed_folder_name}'"


def get_next_label_index() -> int:
    """
    Reads the next_label_index from config.py.
    Defaults to 1 if config.py or the specific line is not found or parsing fails.
    """
    try:
        if not os.path.exists(CONFIG_FILE_PATH):
            print(f"Warning: {CONFIG_FILE_PATH} not found. Defaulting to index 1.")
            return 1
        with open(CONFIG_FILE_PATH, 'r') as f:
            for line in f:
                if line.startswith("next_label_index ="):
                    try:
                        index = int(line.split('=')[1].strip())
                        return index
                    except (ValueError, IndexError) as e:
                        print(f"Warning: Error parsing next_label_index in {CONFIG_FILE_PATH}: {e}. Defaulting to 1.")
                        return 1
        print(f"Warning: 'next_label_index =' line not found in {CONFIG_FILE_PATH}. Defaulting to index 1.")
        return 1
    except IOError as e:
        print(f"Warning: Could not read {CONFIG_FILE_PATH}: {e}. Defaulting to index 1.")
        return 1

def update_label_index(current_index: int) -> None:
    """
    Updates the next_label_index in config.py.
    If current_index is 20, new_index becomes 1. Otherwise, new_index is current_index + 1.
    Creates config.py if it doesn't exist.
    """
    new_index = 1 if current_index == 20 else current_index + 1
    
    lines = []
    found = False
    
    try:
        if os.path.exists(CONFIG_FILE_PATH):
            with open(CONFIG_FILE_PATH, 'r') as f:
                lines = f.readlines()
        
        new_lines = []
        for line in lines:
            if line.startswith("next_label_index ="):
                new_lines.append(f"next_label_index = {new_index}\n")
                found = True
            else:
                new_lines.append(line)
        
        if not found:
            new_lines.append(f"next_label_index = {new_index}\n")
            
        with open(CONFIG_FILE_PATH, 'w') as f:
            f.writelines(new_lines)
            
    except IOError as e:
        print(f"Error: Could not write to {CONFIG_FILE_PATH}: {e}")
# Ensure newline at EOF

def generate_label_docx(template_path: str, output_path: str, data: dict, label_index: int) -> bool:
    """
    Generates a DOCX file from a template, replacing placeholders with data.

    Args:
        template_path: Path to the DOCX template file.
        output_path: Path to save the generated DOCX file.
        data: Dictionary containing data to fill into placeholders.
        label_index: Integer to identify specific placeholders (e.g., Buyer1, Seller1).

    Returns:
        True if generation was successful, False otherwise.
    """
    try:
        if not os.path.exists(template_path):
            print(f"ERROR: Template file not found at {template_path}")
            return False

        # Copy template to output path to work on a copy
        shutil.copy2(template_path, output_path)
        
        doc = docx.Document(output_path)

        buyer_tag = f"Buyer{label_index}"
        seller_tag = f"Seller{label_index}"
        address_tag = f"Address{label_index}"

        # Retrieve and format data
        # format_name is assumed to be in the same module (processing_logic.py)
        buyer_name = format_name(data.get('BYR1NAM1'), data.get('BYR1NAM2'))
        seller_name = format_name(data.get('SLR1NAM1'), data.get('SLR1NAM2'))
        address = data.get('PROPSTRE', '') # Default to empty string if not found

        # New replacement logic targeting hidden runs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.hidden:
                    current_run_text = run.text.strip() # Strip to handle potential spaces around placeholder
                    if current_run_text == buyer_tag:
                        run.text = buyer_name
                        run.font.hidden = False
                    elif current_run_text == seller_tag:
                        run.text = seller_name
                        run.font.hidden = False
                    elif current_run_text == address_tag:
                        run.text = address
                        run.font.hidden = False
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.hidden:
                                current_run_text = run.text.strip() # Strip to handle potential spaces
                                if current_run_text == buyer_tag:
                                    run.text = buyer_name
                                    run.font.hidden = False
                                elif current_run_text == seller_tag:
                                    run.text = seller_name
                                    run.font.hidden = False
                                elif current_run_text == address_tag:
                                    run.text = address
                                    run.font.hidden = False
        
        doc.save(output_path)
        print(f"Successfully generated label document: {output_path}")
        return True

    except FileNotFoundError: # Specifically for shutil.copy2 if template_path is somehow re-checked or output_path dir invalid
        print(f"ERROR: File not found during copy. Template: {template_path}, Output: {output_path}")
        return False
    except Exception as e:
        print(f"ERROR: An unexpected error occurred while generating DOCX {output_path}: {e}")
        return False